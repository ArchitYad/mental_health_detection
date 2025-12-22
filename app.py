import streamlit as st
from pypdf import PdfReader
from docx import Document
from PIL import Image
import pytesseract
import json

# Updated LangChain imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document as LCDocument

# LCEL & Groq
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_groq import ChatGroq

# ----------------------------
# Streamlit Setup
# ----------------------------
st.set_page_config(page_title="Storybook AI Demo", layout="wide")
st.title("📘 Storybook AI Demo with LLaMA 3.1 + Chroma + LCEL")

# Use Streamlit secrets for API key
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "YOUR_GROQ_API_KEY")
if GROQ_API_KEY == "YOUR_GROQ_API_KEY":
    st.error("❌ Set your GROQ_API_KEY in Streamlit secrets or replace placeholder")
    st.stop()

# Groq LLM (valid model)
llm = ChatGroq(
    groq_api_key=GROQ_API_KEY,
    model_name="llama-3.1-8b-instant"
)

# ----------------------------
# Helper Functions
# ----------------------------
def extract_text(file):
    if file.type == "application/pdf":
        reader = PdfReader(file)
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    elif file.type == "text/plain":
        return file.read().decode("utf-8")
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif file.type.startswith("image"):
        image = Image.open(file)
        return pytesseract.image_to_string(image)
    return ""

def chunk_text(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
    return splitter.split_text(text)

# ----------------------------
# Streamlit UI
# ----------------------------
topic = st.text_input("Enter a topic:", "AWS API Gateway")
uploaded_file = st.file_uploader(
    "Upload PDF / TXT / DOCX / Image (optional)",
    type=["pdf","txt","docx","png","jpg","jpeg"]
)

if st.button("Generate Storybook"):

    if not topic and not uploaded_file:
        st.warning("Enter a topic or upload a document")
        st.stop()

    with st.spinner("Generating storybook..."):

        # ---- Step 1: Text Extraction ----
        text = topic
        if uploaded_file:
            text += "\n" + extract_text(uploaded_file)
        st.success("✅ Text extracted")

        # ---- Step 2: Chunking ----
        chunks = chunk_text(text)
        st.info(f"Created {len(chunks)} chunks")

        # ---- Step 3: Embeddings + Chroma ----
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        docs = [LCDocument(page_content=c) for c in chunks]
        vectordb = Chroma.from_documents(docs, embedding=embeddings, persist_directory="./chroma_db")

        # ---- Step 4: LCEL Prompt ----
        template = """
You are a storybook generator. Using the context below,
create a 3-page storybook for the topic: {topic}.
Page 1: Simple explanation
Page 2: Real-world example
Page 3: Technical diagram nodes & edges

Context: {context}

Return strictly valid JSON: {{"pages": [{{"page": 1, "title": "...", "story": "..."}}, ...], "title": "..."}}"""
        
        prompt = PromptTemplate(
            input_variables=["context", "topic"],
            template=template
        )

        # ---- Step 5: Modern LCEL Runnable Chain ----
        chain = (
            {
                "context": vectordb.as_retriever(search_kwargs={"k": 3}) | (lambda docs: "\n\n".join([d.page_content for d in docs])),
                "topic": RunnablePassthrough()
            }
            | prompt
            | llm
            | JsonOutputParser()
        )

        # ---- Step 6: Run Chain ----
        try:
            storybook = chain.invoke(topic)
            st.session_state.storybook = storybook
        except Exception as e:
            st.error(f"Error generating storybook: {e}")
            st.stop()

    # ---- Step 7: Display Storybook ----
    st.header("📖 Generated Storybook")
    storybook = st.session_state.get("storybook", {})

    if "pages" in storybook:
        for page in storybook["pages"]:
            st.subheader(f"Page {page.get('page', '?')}: {page.get('title', 'Untitled')}")
            st.write(page.get("story", ""))
            if "diagram" in page:
                st.code(page["diagram"])
    else:
        st.json(storybook)  # fallback debug view

    if "title" in storybook:
        st.caption(f"**Title:** {storybook['title']}")

    st.success("✅ Storybook generated successfully!")
